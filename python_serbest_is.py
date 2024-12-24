import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import seaborn as sns

# 1. Məlumatları hazırlayın
vulnerabilities = [
    {"ID": 1, "Name": "SQL Injection", "Severity": "High", "Status": "Unresolved"},
    {"ID": 2, "Name": "Cross-Site Scripting (XSS)", "Severity": "Medium", "Status": "Resolved"},
    {"ID": 3, "Name": "Insecure Direct Object Reference", "Severity": "High", "Status": "Unresolved"},
    {"ID": 4, "Name": "Security Misconfiguration", "Severity": "Low", "Status": "Resolved"},
]

df = pd.DataFrame(vulnerabilities)

# 2. Statistik hesabat yaratmaq
severity_count = df['Severity'].value_counts()
status_count = df['Status'].value_counts()

# 3. Vizualizasiya yaradın
sns.set(style="whitegrid")
plt.figure(figsize=(8, 5))
sns.countplot(x='Severity', data=df, order=['Low', 'Medium', 'High'], palette='coolwarm')
plt.title("Vulnerability Severity Distribution")
plt.savefig("severity_distribution.png")  # Şəkli saxlamaq

# 4. Word hesabat yaradın
doc = Document()
doc.add_heading("Penetration Test Report", level=1)

# Məlumat cədvəli əlavə edin
doc.add_heading("Vulnerability Summary", level=2)
table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'

# Cədvəl başlıqları
for idx, col_name in enumerate(df.columns):
    table.cell(0, idx).text = col_name

# Cədvəl məlumatları
for row in df.itertuples(index=False):
    row_cells = table.add_row().cells
    for idx, value in enumerate(row):
        row_cells[idx].text = str(value)

# Statistik məlumatları əlavə edin
doc.add_heading("Statistics", level=2)
doc.add_paragraph(f"Severity Counts:\n{severity_count.to_string()}")
doc.add_paragraph(f"Status Counts:\n{status_count.to_string()}")

# Şəkil əlavə edin
doc.add_heading("Visualization", level=2)
doc.add_picture("severity_distribution.png", width=Inches(5.5))

# Hesabatı saxlayın
doc.save("penetration_test_report.docx")

print("Hesabat yaradıldı: penetration_test_report.docx")
